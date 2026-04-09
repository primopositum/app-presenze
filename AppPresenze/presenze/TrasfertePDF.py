from __future__ import annotations

import copy
import hashlib
import uuid
from dataclasses import dataclass
from datetime import date, timedelta
from decimal import Decimal, ROUND_HALF_UP
from io import BytesIO
from pathlib import Path
import subprocess
import tempfile
from typing import Dict, List, Tuple

from django.conf import settings
from django.http import FileResponse, Http404
from django.utils import timezone

from rest_framework import status
from rest_framework.exceptions import ValidationError
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework.views import APIView

from docx import Document
from docx.shared import Cm

from .models import Contratto, Signature, SignatureEvent, Spesa, Trasferta, Utente


DOCX_TEMPLATE_PATH = Path(
    getattr(
        settings,
        "TRASFERTE_DOCX_TEMPLATE_PATH",
        settings.BASE_DIR / "presenze" / "templates" / "Template_Trasferte.docx",
    )
)

SIGNATURE_WIDTH_CM = 5.0
SIGNATURE_HEIGHT_CM = 2.5


def _to_bool(value: str | None, default: bool = False) -> bool:
    if value is None:
        return default
    return str(value).strip().lower() in {"1", "true", "yes", "on", "si"}


def _italian_month_year(d: date) -> str:
    months = [
        "",
        "Gennaio",
        "Febbraio",
        "Marzo",
        "Aprile",
        "Maggio",
        "Giugno",
        "Luglio",
        "Agosto",
        "Settembre",
        "Ottobre",
        "Novembre",
        "Dicembre",
    ]
    return f"{months[d.month]} {d.year}"


def _format_money(value: Decimal) -> str:
    return f"{value.quantize(Decimal('0.01'), rounding=ROUND_HALF_UP):.2f}".replace(".", ",")


def _format_decimal(value: Decimal) -> str:
    normalized = format(value, "f").rstrip("0").rstrip(".")
    if not normalized:
        normalized = "0"
    return normalized.replace(".", ",")


def _previous_month_range_from(base_date: date) -> Tuple[date, date]:
    first_of_month = base_date.replace(day=1)
    last_prev = first_of_month - timedelta(days=1)
    start_prev = last_prev.replace(day=1)
    return start_prev, last_prev


def _replace_text_in_paragraph(paragraph, replacements: Dict[str, str]) -> None:
    text = "".join(run.text for run in paragraph.runs)
    for key, value in replacements.items():
        text = text.replace(key, value)

    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""


def _replace_text_in_row(row, replacements: Dict[str, str]) -> None:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            _replace_text_in_paragraph(paragraph, replacements)


def _replace_text_in_part(part, replacements: Dict[str, str]) -> None:
    for paragraph in part.paragraphs:
        _replace_text_in_paragraph(paragraph, replacements)
    for table in part.tables:
        for row in table.rows:
            _replace_text_in_row(row, replacements)


def _iter_header_footer_parts(doc: Document):
    for section in doc.sections:
        for attr in (
            "header",
            "first_page_header",
            "even_page_header",
            "footer",
            "first_page_footer",
            "even_page_footer",
        ):
            part = getattr(section, attr, None)
            if part is not None:
                yield part


def _row_contains(row, token: str) -> bool:
    return token in " ".join(cell.text for cell in row.cells)


def _duplicate_row(table, original_row):
    new_row_xml = copy.deepcopy(original_row._tr)
    original_row._tr.addnext(new_row_xml)
    for row in table.rows:
        if row._tr is new_row_xml:
            return row
    return None


def _replace_firma_with_image(doc: Document, firma_token: str, image_path: Path) -> None:
    if not image_path.exists():
        raise ValidationError({"errors": f"File firma non trovato: {image_path}"})

    def process_paragraph(paragraph):
        full_text = "".join(run.text for run in paragraph.runs)
        if firma_token not in full_text:
            return

        for run in paragraph.runs:
            run.text = run.text.replace(firma_token, "")

        img_run = paragraph.add_run()
        img_run.add_picture(
            str(image_path),
            width=Cm(SIGNATURE_WIDTH_CM),
            height=Cm(SIGNATURE_HEIGHT_CM),
        )

    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)

    for part in _iter_header_footer_parts(doc):
        for paragraph in part.paragraphs:
            process_paragraph(paragraph)
        for table in part.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        process_paragraph(paragraph)


def _month_totals_init() -> Dict[int, Decimal]:
    return {
        Spesa.TrasfertaType.PARCHEGGI: Decimal("0.00"),
        Spesa.TrasfertaType.RISTORANTI: Decimal("0.00"),
        Spesa.TrasfertaType.HOTEL: Decimal("0.00"),
        Spesa.TrasfertaType.ALTRO: Decimal("0.00"),
        Spesa.TrasfertaType.PEDAGGI: Decimal("0.00"),
        Spesa.TrasfertaType.KM: Decimal("0.00"),
    }


@dataclass
class TrasfertaRow:
    date: str
    tragitto: str
    km: Decimal
    calc: Decimal
    park: Decimal
    rist: Decimal
    hote: Decimal
    other: Decimal
    ped: Decimal

    @property
    def total(self) -> Decimal:
        return self.calc + self.park + self.rist + self.hote + self.other + self.ped

    def to_placeholders(self) -> Dict[str, str]:
        return {
            "#Date": self.date,
            "#tragitto": self.tragitto,
            "#Azienda": self.tragitto,
            "#Km": _format_decimal(self.km) if self.km > 0 else "",
            "#Calc": _format_money(self.calc) if self.calc > 0 else "",
            "#Park": _format_money(self.park) if self.park > 0 else "",
            "#Rist": _format_money(self.rist) if self.rist > 0 else "",
            "#Hote": _format_money(self.hote) if self.hote > 0 else "",
            "#Other": _format_money(self.other) if self.other > 0 else "",
            "#Ped": _format_money(self.ped) if self.ped > 0 else "",
            "#Tot": _format_money(self.total) if self.total > 0 else "",
        }


def _build_rows_and_totals(trasferte: List[Trasferta]) -> Tuple[List[TrasfertaRow], Dict[int, Decimal], Decimal]:
    rows: List[TrasfertaRow] = []
    month_totals = _month_totals_init()
    grand_total = Decimal("0.00")

    for tr in trasferte:
        grouped: Dict[int, Decimal] = {}
        for spesa in tr.spese.all():
            grouped.setdefault(spesa.type, Decimal("0.00"))
            grouped[spesa.type] += Decimal(str(spesa.importo))

        km_total = grouped.get(Spesa.TrasfertaType.KM, Decimal("0.00"))
        coeff = Decimal("0.00")
        if tr.automobile:
            coeff = Decimal(str(tr.automobile.coefficiente or Decimal("0.00")))
        calc_total = (km_total * coeff).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
        tragitto_value = " / ".join([str(p).strip() for p in (tr.tragitto or []) if str(p).strip()])

        row = TrasfertaRow(
            date=tr.data.strftime("%d/%m/%Y"),
            tragitto=tragitto_value,
            km=km_total,
            calc=calc_total,
            park=grouped.get(Spesa.TrasfertaType.PARCHEGGI, Decimal("0.00")),
            rist=grouped.get(Spesa.TrasfertaType.RISTORANTI, Decimal("0.00")),
            hote=grouped.get(Spesa.TrasfertaType.HOTEL, Decimal("0.00")),
            other=(
                grouped.get(Spesa.TrasfertaType.BIGLIETTI, Decimal("0.00"))
                + grouped.get(Spesa.TrasfertaType.ALTRO, Decimal("0.00"))
            ),
            ped=grouped.get(Spesa.TrasfertaType.PEDAGGI, Decimal("0.00")),
        )
        rows.append(row)

        month_totals[Spesa.TrasfertaType.PARCHEGGI] += row.park
        month_totals[Spesa.TrasfertaType.RISTORANTI] += row.rist
        month_totals[Spesa.TrasfertaType.HOTEL] += row.hote
        month_totals[Spesa.TrasfertaType.ALTRO] += row.other
        month_totals[Spesa.TrasfertaType.PEDAGGI] += row.ped
        month_totals[Spesa.TrasfertaType.KM] += row.calc

        grand_total += row.total

    return rows, month_totals, grand_total


def _fill_docx(
    template_path: Path,
    general_data: Dict[str, str],
    rows: List[TrasfertaRow],
    totals_map: Dict[str, str],
    firma: bool,
    firma_path: Path | None,
) -> BytesIO:
    if not template_path.exists():
        raise Http404(f"Template DOCX non trovato: {template_path}")

    doc = Document(str(template_path))
    merged_general = {**general_data, **totals_map}

    _replace_text_in_part(doc, merged_general)
    for part in _iter_header_footer_parts(doc):
        _replace_text_in_part(part, merged_general)

    for table in doc.tables:
        template_row = None
        for row in table.rows:
            if _row_contains(row, "#Date"):
                template_row = row
                break

        if template_row is None:
            for row in table.rows:
                _replace_text_in_row(row, merged_general)
            continue

        row_objs = [template_row]
        for _ in rows[1:]:
            new_row = _duplicate_row(table, row_objs[-1])
            if new_row is None:
                raise ValidationError({"errors": "Impossibile duplicare la riga template nel documento."})
            row_objs.append(new_row)

        if rows:
            for row_obj, data in zip(row_objs, rows):
                _replace_text_in_row(row_obj, data.to_placeholders())
        else:
            _replace_text_in_row(row_objs[0], TrasfertaRow("", "", Decimal("0"), Decimal("0"), Decimal("0"), Decimal("0"), Decimal("0"), Decimal("0"), Decimal("0")).to_placeholders())

        for row in table.rows:
            if row not in row_objs:
                _replace_text_in_row(row, merged_general)

    if firma:
        if not firma_path:
            raise ValidationError({"errors": "firma=true richiede 'firma_path'."})
        _replace_firma_with_image(doc, "#FIRMA", firma_path)
    else:
        _replace_text_in_part(doc, {"#FIRMA": ""})
        for part in _iter_header_footer_parts(doc):
            _replace_text_in_part(part, {"#FIRMA": ""})

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def _convert_docx_to_pdf(docx_bytes: bytes) -> BytesIO:
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp = Path(tmp_dir)
        docx_path = tmp / "trasferte.docx"
        pdf_path = tmp / "trasferte.pdf"
        docx_path.write_bytes(docx_bytes)

        conversion_errors: List[str] = []

        # Conversione DOCX -> PDF via LibreOffice headless.
        try:
            subprocess.run(
                [
                    "soffice",
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(tmp),
                    str(docx_path),
                ],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
            )
        except FileNotFoundError:
            conversion_errors.append("soffice non trovato nel container")
        except subprocess.CalledProcessError as exc:
            stderr = (exc.stderr or "").strip()
            stdout = (exc.stdout or "").strip()
            details = stderr or stdout or str(exc)
            conversion_errors.append(f"soffice: {details}")
        except Exception as exc:
            conversion_errors.append(f"soffice: {exc}")

        if not pdf_path.exists():
            raise ValidationError(
                {
                    "errors": (
                        "Conversione PDF non disponibile via LibreOffice headless. "
                        f"Dettagli: {' | '.join(conversion_errors)}"
                    )
                }
            )

        out = BytesIO(pdf_path.read_bytes())
        out.seek(0)
        return out


def _get_client_ip(request) -> str | None:
    xff = request.META.get("HTTP_X_FORWARDED_FOR")
    if xff:
        return xff.split(",")[0].strip()
    return request.META.get("REMOTE_ADDR")


def _signature_to_temp_file(signature: Signature, temp_dir_path: str) -> Path | None:
    mime_ext = {
        "image/png": "png",
        "image/jpeg": "jpg",
        "image/jpg": "jpg",
        "image/webp": "webp",
        "image/bmp": "bmp",
        "image/gif": "gif",
        "image/svg+xml": "svg",
    }
    mime = (signature.mime_type or "").lower().strip()
    ext = mime_ext.get(mime, "png")

    out_path = Path(temp_dir_path) / f"signature_{signature.id}.{ext}"
    if signature.image_data:
        out_path.write_bytes(bytes(signature.image_data))
        return out_path
    if signature.svg:
        out_path = Path(temp_dir_path) / f"signature_{signature.id}.svg"
        out_path.write_text(signature.svg, encoding="utf-8")
        return out_path
    return None


def build_trasferte_pdf_bytes(
    user: Utente,
    trasferte: List[Trasferta],
    period_start: date,
    *,
    firma: bool = False,
    firma_path: Path | None = None,
    reference_trasferta: Trasferta | None = None,
) -> bytes:
    """
    Costruisce il PDF trasferte a partire da una lista di trasferte già filtrata.
    """
    rows, month_totals, grand_total = _build_rows_and_totals(trasferte)

    nome = getattr(user, "nome", "") or ""
    cognome = getattr(user, "cognome", "") or ""
    full_name = f"{nome} {cognome}".strip() or user.email
    contract = (
        Contratto.objects
        .filter(utente_id=user.id, is_active=True)
        .order_by("-data_ass", "-id")
        .first()
    )
    role = (contract.tipologia if contract else "") or ""
    auto_source = reference_trasferta
    if auto_source is None and trasferte:
        auto_source = trasferte[0]

    auto_descr = ""
    rimborso = "0"
    if auto_source and auto_source.automobile:
        auto_descr = auto_source.automobile.descrizione or str(auto_source.automobile)
        rimborso = _format_decimal(Decimal(str(auto_source.automobile.coefficiente or Decimal("0.00"))))

    general_data = {
        "#MeseAnno": _italian_month_year(period_start),
        "#NomeCognome": full_name,
        "#Role" : role,
        "#Autovettura": auto_descr,
        "#Rimborso": rimborso,
        "#data_audit": timezone.localdate().isoformat(),
    }

    totals_map = {
        "#1": _format_money(month_totals[Spesa.TrasfertaType.KM]) if month_totals[Spesa.TrasfertaType.KM] > 0 else "",
        "#2": _format_money(month_totals[Spesa.TrasfertaType.PARCHEGGI]) if month_totals[Spesa.TrasfertaType.PARCHEGGI] > 0 else "",
        "#3": _format_money(month_totals[Spesa.TrasfertaType.RISTORANTI]) if month_totals[Spesa.TrasfertaType.RISTORANTI] > 0 else "",
        "#4": _format_money(month_totals[Spesa.TrasfertaType.HOTEL]) if month_totals[Spesa.TrasfertaType.HOTEL] > 0 else "",
        "#5": _format_money(month_totals[Spesa.TrasfertaType.PEDAGGI]) if month_totals[Spesa.TrasfertaType.PEDAGGI] > 0 else "",
        "#6": _format_money(month_totals[Spesa.TrasfertaType.ALTRO]) if month_totals[Spesa.TrasfertaType.ALTRO] > 0 else "",
        "#7": _format_money(grand_total) if grand_total > 0 else "",
    }

    try:
        docx_stream = _fill_docx(
            template_path=DOCX_TEMPLATE_PATH,
            general_data=general_data,
            rows=rows,
            totals_map=totals_map,
            firma=firma,
            firma_path=firma_path,
        )
    except (Http404, ValidationError):
        raise
    except Exception as exc:
        raise ValidationError({"errors": f"Errore compilazione documento: {exc}"}) from exc

    return _convert_docx_to_pdf(docx_stream.getvalue()).getvalue()


class TrasfertePDFView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        data_str = request.query_params.get("data")
        utente_id_param = request.query_params.get("u_id")
        firma = _to_bool(request.query_params.get("firma"), default=False)
        firma_path_param = request.query_params.get("firma_path")
        firma_status = "ok"

        if not utente_id_param:
            return Response(
                {"errors": "Parametro 'u_id' obbligatorio."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        try:
            utente_id = int(utente_id_param)
        except (TypeError, ValueError):
            return Response(
                {"errors": "Parametro 'u_id' non valido."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        base_date = timezone.localdate()
        if data_str:
            try:
                base_date = date.fromisoformat(data_str)
            except ValueError:
                return Response(
                    {"errors": "Parametro 'data' non valido. Usa YYYY-MM-DD."},
                    status=status.HTTP_400_BAD_REQUEST,
                )

        is_super = request.user.is_superuser
        is_owner = request.user.id == utente_id
        if not (is_super or is_owner):
            return Response(
                {"errors": "Non hai i permessi per questo utente."},
                status=status.HTTP_403_FORBIDDEN,
            )

        try:
            user = Utente.objects.get(pk=utente_id)
        except Utente.DoesNotExist:
            return Response(
                {"errors": "Utente non trovato."},
                status=status.HTTP_404_NOT_FOUND,
            )

        start_prev_month, end_prev_month = _previous_month_range_from(base_date)

        trasferte = list(
            Trasferta.objects.filter(
                utente_id=user.id,
                data__range=(start_prev_month, end_prev_month),
            )
            .select_related("automobile")
            .prefetch_related("spese")
            .order_by("data", "id")
        )

        firma_path = None
        signature_used = None
        temp_dir = None
        if firma:
            if firma_path_param:
                firma_path = Path(firma_path_param)
            else:
                signature_used = (
                    Signature.objects
                    .filter(user_id=user.id)
                    .order_by("-created_at")
                    .first()
                )
                if signature_used is None:
                    firma = False
                    firma_status = "firma mancante"
                temp_dir = tempfile.TemporaryDirectory()
                if signature_used is not None:
                    firma_path = _signature_to_temp_file(signature_used, temp_dir.name)
                    if firma_path is None:
                        firma = False
                        firma_status = "firma mancante"

        try:
            pdf_bytes = build_trasferte_pdf_bytes(
                user=user,
                trasferte=trasferte,
                period_start=start_prev_month,
                firma=firma,
                firma_path=firma_path,
            )
        finally:
            if temp_dir is not None:
                temp_dir.cleanup()

        if firma and signature_used is not None:
            SignatureEvent.objects.create(
                signature=signature_used,
                user=request.user,
                event_type=SignatureEvent.EventType.USED,
                document_id=uuid.uuid4(),
                document_sha256=hashlib.sha256(pdf_bytes).hexdigest(),
                ip_address=_get_client_ip(request),
                user_agent=request.META.get("HTTP_USER_AGENT", ""),
            )

        pdf_stream = BytesIO(pdf_bytes)
        filename = f"trasferte_{user.id}_{start_prev_month.strftime('%Y_%m')}.pdf"
        response = FileResponse(pdf_stream, as_attachment=True, filename=filename)
        response["X-Firma-Status"] = firma_status
        response["Access-Control-Expose-Headers"] = "X-Firma-Status"
        return response

